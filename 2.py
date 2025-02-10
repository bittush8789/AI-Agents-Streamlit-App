import streamlit as st
from crewai import Agent, Task, Crew
from langchain_openai import ChatOpenAI
from docx import Document
from io import BytesIO
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()


# LLM object and API Key
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")

# Configure the page
st.set_page_config(
    page_title="AI Agents to Empower Doctors",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Add custom CSS
st.markdown("""
    <style>
    .stApp {
        max-width: 1200px;
        margin: 0 auto;
    }
    .stButton>button {
        width: 100%;
    }
    .diagnosis-result {
        background-color: white;
        padding: 20px;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    </style>
""", unsafe_allow_html=True)

def generate_docx(result):
    doc = Document()
    doc.add_heading('Healthcare Diagnosis and Treatment Recommendations', 0)
    doc.add_paragraph(result)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def main():
    st.title("AI Agents to Empower Doctors")
    
    # Create form
    with st.form("patient_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            gender = st.selectbox(
                'Gender',
                ('Male', 'Female', 'Other')
            )
            age = st.number_input('Age', min_value=0, max_value=120, value=25)
        
        with col2:
            symptoms = st.text_area('Symptoms', placeholder='e.g., fever, cough, headache')
            medical_history = st.text_area('Medical History', placeholder='e.g., diabetes, hypertension')
        
        submit_button = st.form_submit_button("Get Diagnosis and Treatment Plan")

    if submit_button:
        if not symptoms:
            st.error("Please enter symptoms before submitting.")
            return

        with st.spinner('Generating recommendations...'):
            try:
                # Initialize LLM
                llm = ChatOpenAI(
                    model="gpt-3.5-turbo-16k",
                    temperature=0.1,
                    max_tokens=8000
                )

                # Create Agents
                diagnostician = Agent(
                    role="Medical Diagnostician",
                    goal="Analyze patient symptoms and medical history to provide a preliminary diagnosis.",
                    backstory="Expert in diagnosing medical conditions based on symptoms and history.",
                    verbose=True,
                    allow_delegation=False,
                    llm=llm
                )

                treatment_advisor = Agent(
                    role="Treatment Advisor",
                    goal="Recommend appropriate treatment plans based on diagnosis.",
                    backstory="Specialist in creating personalized treatment plans.",
                    verbose=True,
                    allow_delegation=False,
                    llm=llm
                )

                # Create Tasks
                diagnose_task = Task(
                    description=f"Analyze symptoms: {symptoms}\nMedical history: {medical_history}\nAge: {age}\nGender: {gender}\nProvide preliminary diagnosis.",
                    agent=diagnostician
                )

                treatment_task = Task(
                    description="Based on the diagnosis, recommend a detailed treatment plan.",
                    agent=treatment_advisor
                )

                # Create and run Crew
                crew = Crew(
                    agents=[diagnostician, treatment_advisor],
                    tasks=[diagnose_task, treatment_task],
                    verbose=2
                )

                result = crew.kickoff()

                # Display results
                st.markdown("### Results")
                st.markdown('<div class="diagnosis-result">', unsafe_allow_html=True)
                st.write(result)
                st.markdown('</div>', unsafe_allow_html=True)

                # Generate and offer document download
                doc = generate_docx(result)
                st.download_button(
                    label="Download Report",
                    data=doc,
                    file_name="diagnosis_and_treatment_plan.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            except Exception as e:
                st.error(f"An error occurred: {str(e)}")

if __name__ == "__main__":
    main()